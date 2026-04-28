from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
OUTPUT_PATH = BASE_DIR / "AI_Resume_Analyzer_IEEE_Report.docx"


def set_page_margins(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)


def set_columns(section, count):
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    cols_el = cols[0] if cols else OxmlElement("w:cols")
    cols_el.set(qn("w:num"), str(count))
    cols_el.set(qn("w:space"), "360")
    if not cols:
        sect_pr.append(cols_el)


def set_default_style(document):
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(10)


def add_title(document, text, subtitle_lines):
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(20)

    for line in subtitle_lines:
        para = document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(line)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        run.font.size = Pt(10)


def add_abstract(document, abstract_text, keywords_text):
    abstract = document.add_paragraph()
    abstract.paragraph_format.space_before = Pt(6)
    abstract.paragraph_format.space_after = Pt(6)
    abstract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    label = abstract.add_run("Abstract—")
    label.bold = True
    label.italic = True
    abstract.add_run(abstract_text)

    keywords = document.add_paragraph()
    keywords.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    label = keywords.add_run("Index Terms—")
    label.bold = True
    label.italic = True
    keywords.add_run(keywords_text)


def add_heading(document, text):
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(5)
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(10)


def add_subheading(document, text):
    para = document.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(10)


def add_body_paragraph(document, text):
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.first_line_indent = Inches(0.17)
    para.paragraph_format.line_spacing = 1.0
    para.paragraph_format.space_after = Pt(4)
    para.add_run(text)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_placeholder_box(document, title_text, caption_text, height_inches=2.45):
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run(title_text)
    run.bold = True

    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    row = table.rows[0]
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    row.height = Inches(height_inches)
    cell = row.cells[0]
    cell.width = Inches(6.8)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    shade_cell(cell, "F4F6F8")

    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after = Pt(6)
    label = para.add_run("[ Insert Screenshot Here ]")
    label.bold = True
    label.font.size = Pt(12)
    label.font.color.rgb = RGBColor(80, 87, 99)

    note = cell.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note.add_run(caption_text)
    note_run.italic = True
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(100, 116, 139)


def add_references(document, references):
    add_heading(document, "REFERENCES")
    for ref in references:
        para = document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.first_line_indent = Inches(-0.25)
        para.paragraph_format.space_after = Pt(2)
        para.add_run(ref)


def build_report():
    document = Document()
    set_default_style(document)

    first_section = document.sections[0]
    set_page_margins(first_section)
    set_columns(first_section, 1)

    add_title(
        document,
        "AI Resume Analyzer: Design and Implementation of a Dataset-Driven Resume Screening System",
        [
            "Author Name: ________________________________",
            "Affiliation / Department: ________________________________",
            "Project Report Prepared in IEEE-Style Word Format",
        ],
    )

    add_abstract(
        document,
        " This report presents the design and implementation of an AI Resume Analyzer developed as a "
        "modern web application using Flask, HTML, CSS, and JavaScript. The system accepts PDF resumes, "
        "extracts textual content, evaluates resume completeness, identifies education and experience "
        "signals, matches role-specific skills, and produces a score out of one hundred. To improve "
        "practical accuracy, the final system integrates a structured technology-jobs dataset, supports "
        "custom skills and custom keywords, applies synonym-aware matching, and uses TF-IDF similarity "
        "through scikit-learn to estimate role relevance. The application also generates downloadable "
        "reports and offers a responsive interface designed for real-world screening demonstrations.",
        "resume analysis, ATS optimization, Flask web application, PDF text extraction, TF-IDF similarity, "
        "job-role matching, technology jobs dataset, keyword optimization",
    )

    body_section = document.add_section(WD_SECTION.CONTINUOUS)
    set_page_margins(body_section)
    set_columns(body_section, 2)

    add_heading(document, "I. INTRODUCTION")
    add_body_paragraph(
        document,
        "Recruitment workflows increasingly depend on digital resume screening to reduce manual effort and "
        "to surface candidates whose profiles align with open roles. However, many student-level resume "
        "analyzers rely only on short static keyword lists, which limits adaptability and reduces accuracy "
        "when resumes use alternate phrasing, tool-specific terminology, or broader role language. The AI "
        "Resume Analyzer project was built to address that limitation by combining a clean web interface "
        "with a more structured backend analysis pipeline."
    )
    add_body_paragraph(
        document,
        "The implemented system supports upload of PDF resumes, extraction of resume text, identification "
        "of education and experience details, missing-section checks, role-based skill matching, keyword "
        "gap analysis, improvement suggestions, and report export. In its upgraded form, the project also "
        "uses a local technology-jobs dataset that expands the analyzer beyond three manually defined roles "
        "into a broader set of engineering, data, cloud, quality, and security roles."
    )

    add_heading(document, "II. PROJECT OBJECTIVES AND CONTRIBUTIONS")
    add_body_paragraph(
        document,
        "The primary objective of the project is to provide an intelligent but lightweight resume screening "
        "tool that can be executed locally without dependence on paid APIs. The system aims to help users "
        "understand how well a resume aligns with a selected role, identify missing skills or sections, and "
        "suggest targeted improvements in language, structure, and technical coverage."
    )
    add_body_paragraph(
        document,
        "The major contributions of the final implementation are as follows: a full-stack Flask-based web "
        "application; PDF resume ingestion; dataset-driven role matching for multiple technology jobs; "
        "support for user-defined custom skills and custom keywords; synonym-aware term matching; TF-IDF "
        "similarity scoring; modern responsive user interface design; and generation of a downloadable "
        "analysis report for screening review."
    )

    add_heading(document, "III. SYSTEM ARCHITECTURE AND WORKFLOW")
    add_subheading(document, "A. Frontend Interaction Layer")
    add_body_paragraph(
        document,
        "The frontend is implemented using a single responsive landing page with drag-and-drop resume upload, "
        "a dynamic tech-role selector, optional custom targeting fields, animated score components, and a "
        "dark mode interface. Result cards present overall score, skills found, missing skills, keyword "
        "coverage, role similarity, role market signal, and suggestions. This structure allows the user to "
        "understand both the screening score and the reasoning behind it."
    )
    add_subheading(document, "B. Backend Processing Pipeline")
    add_body_paragraph(
        document,
        "The backend is implemented in Flask with clearly separated helper functions for file validation, "
        "PDF text extraction, section detection, skill matching, keyword coverage calculation, similarity "
        "estimation, scoring, and report generation. The core route accepts the uploaded file and selected "
        "role, extracts the resume text using PyPDF2, performs analysis, and returns structured JSON that "
        "drives the frontend dashboard."
    )
    add_subheading(document, "C. Analysis Workflow")
    add_body_paragraph(
        document,
        "The analysis workflow begins with text normalization and simple section detection for Education, "
        "Experience, Projects, Skills, and Certifications. The system then compares resume text against the "
        "selected role's dataset profile, including core skills, optional skills, role keywords, and related "
        "tools. A final score is produced from role-aligned skills, completeness, keyword coverage, and "
        "role-fit similarity, with small penalties applied when the resume is too short or lacks important "
        "sections."
    )

    add_heading(document, "IV. DATASET-DRIVEN ROLE MODEL")
    add_body_paragraph(
        document,
        "A major enhancement in the project is the introduction of a structured local technology-jobs dataset "
        "stored as JSON. Instead of hardcoding only a few roles, the system now loads multiple curated roles "
        "such as Software Engineer, Frontend Developer, Backend Developer, Full Stack Developer, Data "
        "Scientist, Data Analyst, Machine Learning Engineer, DevOps Engineer, QA Engineer, and Cybersecurity "
        "Analyst. Each role contains category metadata, market signal, outlook score, profile text, must-have "
        "skills, nice-to-have skills, keywords, and related tools."
    )
    add_body_paragraph(
        document,
        "This dataset structure improves maintainability because additional roles can be introduced without "
        "rewriting the analysis engine. It also improves accuracy because each role contains richer semantic "
        "context than a flat keyword list. The analyzer can therefore reason over missing core skills, "
        "adjacent tools, and role-specific language rather than only checking whether a small set of words is present."
    )

    add_heading(document, "V. ACCURACY ENHANCEMENTS")
    add_subheading(document, "A. Synonym-Aware Matching")
    add_body_paragraph(
        document,
        "The analyzer does not rely solely on exact string matches. It maps common technology variations and "
        "aliases such as JS to JavaScript, sklearn to Scikit-learn, CI/CD to continuous integration patterns, "
        "and A/B testing to experimentation language. This makes the system more robust when resumes use real-world "
        "phrasing instead of textbook terminology."
    )
    add_subheading(document, "B. Custom Skills and Keywords")
    add_body_paragraph(
        document,
        "The final interface also accepts user-defined custom skills and custom keywords. This allows the "
        "analyzer to adapt to role-specific openings or job descriptions without editing the source code. "
        "For example, a recruiter or student can add stack-specific terms such as LangChain, Airflow, "
        "Kubernetes, prompt engineering, or deployment automation and immediately evaluate whether the resume reflects them."
    )
    add_subheading(document, "C. TF-IDF Similarity")
    add_body_paragraph(
        document,
        "To improve screening accuracy further, the project uses TF-IDF vectorization through scikit-learn to "
        "estimate similarity between the resume text and the target role profile. This does not replace skill "
        "matching; instead, it complements it by capturing broader role alignment across phrases, tools, and "
        "responsibility language. When scikit-learn is unavailable, the system falls back gracefully to token-overlap scoring."
    )

    add_heading(document, "VI. IMPLEMENTATION DETAILS")
    add_body_paragraph(
        document,
        "The application is organized around a main Flask file, HTML templates, CSS styling, JavaScript "
        "behavior, a requirements file, and a local dataset directory. The primary routes include the "
        "homepage for interaction, the analysis route for resume processing, and a report download route for "
        "export. The frontend submits the uploaded PDF and user-selected role to the backend asynchronously, "
        "which enables a smooth result-rendering experience without full page reloads."
    )
    add_body_paragraph(
        document,
        "The report generation module creates a plain-text downloadable summary of the analysis outcome. "
        "Although lightweight, this reporting layer is useful for demonstrations, screening records, or "
        "future extensions into PDF or Word export. The current implementation also includes a run/stop "
        "guide for local usage and supports persistent server execution in the background."
    )

    add_heading(document, "VII. RESULTS, DISCUSSION, AND LIMITATIONS")
    add_body_paragraph(
        document,
        "The completed system is functional as a local ATS-style analyzer and provides interpretable output "
        "for both technical and academic demonstration settings. The dataset-driven upgrade significantly "
        "improves flexibility because role definitions, related tools, and role signals can be extended "
        "without redesigning the frontend. The TF-IDF layer also provides a more informative estimate of "
        "resume-role fit than strict keyword matching alone."
    )
    add_body_paragraph(
        document,
        "A key limitation is that the project does not yet use a labeled benchmark dataset of resumes and "
        "ground-truth job outcomes. Therefore, the term accuracy in this context refers to stronger practical "
        "matching behavior rather than statistically validated classifier performance. Future work can address "
        "this by adding job-description ingestion, benchmarking on annotated data, and role-specific weighting "
        "for must-have skills, optional skills, and domain terms."
    )

    add_heading(document, "VIII. CONCLUSION")
    add_body_paragraph(
        document,
        "The AI Resume Analyzer project demonstrates how a modern full-stack application can combine "
        "traditional file processing, structured role intelligence, keyword analysis, and lightweight "
        "machine learning to produce a practical resume-screening tool. The final system is visually polished, "
        "easy to run locally, and more realistic than a basic student project because it includes custom "
        "targeting, downloadable reports, expanded technology-role coverage, and dataset-driven analysis logic. "
        "Overall, the project offers a strong foundation for future academic extension or product prototyping."
    )

    add_references(
        document,
        [
            "[1] Pallets Project, \"Flask Documentation,\" Available: https://flask.palletsprojects.com/.",
            "[2] PyPDF2 Project, \"PyPDF2 Documentation,\" Available: https://pypdf2.readthedocs.io/.",
            "[3] F. Pedregosa et al., \"Scikit-learn: Machine Learning in Python,\" Journal of Machine Learning Research, vol. 12, pp. 2825-2830, 2011.",
            "[4] O*NET Resource Center, \"O*NET Database and Technology Skills,\" Available: https://www.onetcenter.org/.",
            "[5] European Commission, \"ESCO: European Skills, Competences, Qualifications and Occupations,\" Available: https://esco.ec.europa.eu/."
        ],
    )

    appendix_section = document.add_section(WD_SECTION.NEW_PAGE)
    set_page_margins(appendix_section)
    set_columns(appendix_section, 1)

    add_heading(document, "APPENDIX: SCREENSHOT PLACEHOLDERS")
    intro = document.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.add_run(
        "The following boxes are intentionally reserved so that final project screenshots can be inserted before submission."
    )

    add_placeholder_box(
        document,
        "Figure A1. Home Page / Upload Interface",
        "Insert a screenshot showing the landing page, drag-and-drop upload area, role selector, and theme toggle."
    )
    add_placeholder_box(
        document,
        "Figure A2. Resume Analysis Dashboard",
        "Insert a screenshot showing the score gauge, keyword coverage, role-fit section, and result cards."
    )
    add_placeholder_box(
        document,
        "Figure A3. Dataset-Driven Role Expansion / New Tech Roles",
        "Insert a screenshot showing the expanded tech-role dropdown or analysis result for a role such as DevOps Engineer or Frontend Developer."
    )

    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_report()
